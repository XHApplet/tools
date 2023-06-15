# -*- coding:utf-8 -*-
"""
python 3.8
pip install python-docx
"""

import os
import docx
import datetime

from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT

filename = "喂养时间表.docx"
dt = datetime.datetime(2023, 6, 12)
page_day = 7    # 一页8天
page_num = 10   # 总共10页
sweek = ["一", "二", "三", "四", "五", "六", "日",]

try:
    os.remove(filename)
except:
    pass

lst_info = [
    ["时间段", "活动", "目标", ],
    ["5-8", "喂奶+玩", "120ml",],
    ["8-10", "睡觉+喂奶", "120ml",],
    ["10-11", "玩一下", "",],
    ["11-12", "喂辅食", "肉蛋轮换",],
    ["12-13", "玩一下", "",],
    ["13-15", "睡觉+喂奶", "120ml",],
    ["15-16", "玩一下", "",],
    ["16-17", "喂辅食", "水果等",],
    ["17-19", "喂奶", "120ml",],
    ["19-20", "玩一下", "",],
    ["20-24", "喂奶+睡觉", "120ml",],
    ["00-05", "睡觉", "120ml",],
]


# 创建一个新的Word文档
doc = docx.Document()

# 设置页面方向为横向,使用了 Enum 类型 WD_ORIENT 来设置页面方向，WD_ORIENT.LANDSCAPE 表示“横向”，WD_ORIENT.PORTRAIT 表示“纵向”。
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width      # 页面宽度
section.page_height = new_height    # 页面高度

section.top_margin = Inches(0.4)  # 顶部边距
section.bottom_margin = Inches(0.4)  # 底部边距
section.left_margin = Inches(0.4)  # 左侧边距
section.right_margin = Inches(0.4)  # 右侧边距

for i in range(page_num):
    st = dt + datetime.timedelta(i * page_day)
    et = st + datetime.timedelta(page_day - 1)
    xst = st.strftime("%Y年%m月%d日")
    xet = et.strftime("%Y年%m月%d日")
    heading = f"喂养记录: {xst} - {xet}"

    # 添加标题
    doc.add_heading(heading, 0)

    lst = []
    for x in range(page_day):
        xday = st + datetime.timedelta(x)
        lst.append(xday.strftime("%m.%d"))

    table = doc.add_table(rows=len(lst_info), cols=len(lst_info[0]) + page_day)
    table.style = 'Table Grid'

    for x, lst_value in enumerate(lst_info):
        row = table.rows[x]
        hdr_cells = row.cells
        row.height = Cm(1.2)
        for j, msg in enumerate(lst_value):
            hdr_cells[j].text = msg
        for k, xx in enumerate(lst):
            msg = "" if x else xx + sweek[k]
            hdr_cells[j + k + 1].text = msg

    if i < page_num - 1:
        doc.add_page_break()

# 保存文档
doc.save(filename)
